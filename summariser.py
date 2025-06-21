"""
High-level orchestration layer in charge of :

1. Lot detection – find the subset of procurement lots that match the
   company’s expertise.

2. Iterative summarisation – chunk large docs, call the LLM, merge partial
   answers, then deduplicate.

3. Markdown → DOCX helpers – keep rich formatting when injecting markdown
   into python-docx.

4. Final assembly – build the Word.docx synthesis (text + images).

The public surface
------------------
``set_llm`` / ``extract_interesting_lots`` / ``summarize_documents`` /
``final_summary_and_export`` – those four are used by *app.py*.
"""

import json
import os
import streamlit as st
from typing import Dict, List, Callable

from markdown import markdown
from bs4 import BeautifulSoup, NavigableString
from docx.shared import Pt, Inches
from docx import Document as DocxDocument
from PIL import Image
from langchain.schema import HumanMessage, BaseMessage          

from token_utils import count_tokens, split_text
from loaders      import load_pdf, load_docx, load_excel
from lexicon      import normalize, rerank
from image_utils  import convert_image_bytes
from vectorstore_builder import generate_image_descriptions
from config       import CLIENT, CLIENT_DESCRIPTION

# LLM injection--------------------------------------------------------------
llm: Callable[[List[BaseMessage]], object] | None = None

def set_llm(llm_client: Callable[[List[BaseMessage]], object]) -> None:
    """
    Register the LLM call-able that all helper routines will use.

    Must be called once at application start-up *after* you have created
    the ChatMistralAI / ChatOpenAI instance.
    """
    global llm
    llm = llm_client

# Markdown → python-docx-----------------------------------------------------
def add_formatted_text(paragraph,element,bold: bool = False,italic: bool = False,inline: bool = False,) -> None:
    """
    Recursively copy element (BeautifulSoup node) into a python-docx
    paragraph while preserving bold / italic.

    Parameters
    ----------
    paragraph : docx.text.paragraph.Paragraph
        Destination paragraph.
    element : bs4.element.Tag | str
        Current node in the HTML tree.
    bold, italic : bool
        State propagated through recursion.
    inline : bool
        If True, nested ``<p>`` tags are treated as soft line-breaks instead
        of creating a brand-new paragraph.
    """

    if isinstance(element, NavigableString):
        text = str(element)
        if inline:
            text = text.replace('\n', ' ')
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        return

    if element.name == "br":
        paragraph.add_run().add_break()
        return

    for child in element.children:

        if isinstance(child, NavigableString):
            text = str(child)
            
            if inline:
                text = text.replace('\n', ' ')

            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic

        else:
            child_bold = bold or (child.name in ["strong", "b"])
            child_italic = italic or (child.name in ["em", "i"])
            
            if inline and child.name in ["p"]:
                add_formatted_text(paragraph, child, bold=child_bold, italic=child_italic, inline=True)

            else:
                add_formatted_text(paragraph, child, bold=child_bold, italic=child_italic, inline=inline)

# Lot extraction (LLM-powered)----------------------------------------------
def extract_interesting_lots(lots_file_path):
    """
    Ask the LLM to identify which procurement lots are relevant to the
    company, based on the “lots” document uploaded by the user.

    The full text of the file (PDF/DOCX/XLS) is inlined into a prompt
    that returns only a Python-dict-like string.

    Returns
    -------
    str
        Raw LLM response (dict-as-string).  Caller post-processes it.
    """
    # --- Load the document into plain text ----------------------------------
    if lots_file_path.lower().endswith(".pdf"):
        lots_text = load_pdf(lots_file_path)

    elif lots_file_path.lower().endswith(".docx"):
        lots_text = load_docx(lots_file_path)

    elif lots_file_path.lower().endswith((".xlsx", ".xls")):
        lots_text = load_excel(lots_file_path)

    else:
        lots_text = ""

    # --- Prompt the LLM -----------------------------------------------------
    prompt_template = """
    Tu es un expert en extraction d'informations. 
    Ton but est d'extraire les lots pertinents en fonction de la mission de l'entreprise, il peut y en avoir un, ou eventuellement deux. 
    Il est impératif que ces lots correspondent précisément à la mission de l'entreprise. 
    Retourne la liste des lots pertinents sous forme d'un dictionnaire python, avec comme clé le Numéro/Nom du lot, et comme value l'ensemble des informations relatives au lot que tu as pu trouver. 
    Ne retourne que le dictionnaire et rien d'autre.
    
    La mission de l'entreprise est la suivante : {client_description}
    
    Ne retourne que le dictionnaire et rien d'autre. Voici le document contenant les lots : {content}
    """

    prompt = prompt_template.format(content=lots_text,client_description=CLIENT_DESCRIPTION)
    response = llm([HumanMessage(content=prompt)])
    return response.content.strip()

# Summarise every document uploaded-------------------------------------------
def summarize_documents(pdf_folder: str, interesting_lots: str) -> Dict[str, str]:
    """
    Iterate over all files in pdf_folder, extract / OCR each one and let the
    LLM generate two summaries per fragment:

    * details specific to the target lot(s)
    * general chantier information

    The two answers are concatenated and stored under the file’s name.

    Returns
    -------
    dict[str, str]
        ``{"file.pdf": "Informations sur le lot : …"}``
    """

    summaries   = {}

    files       = os.listdir(pdf_folder)
    total       = len(files)

    for idx, file in enumerate(files, 1):
        print(f"{round(idx/total*100)} %  {file}")

        if file.lower().endswith((".pdf", ".docx", ".xlsx", ".xls")):
            file_path = os.path.join(pdf_folder, file)
        
            # --- Text extraction ------------------------------------------------
            if file.lower().endswith(".pdf"):
                text = load_pdf(file_path)

            elif file.lower().endswith(".docx"):
                text = load_docx(file_path)

            elif file.lower().endswith((".xlsx", ".xls")):
                text = load_excel(file_path)
            else:
                continue

            # --- Split long docs into token-bounded chunks ---------------------
            if count_tokens(text, encoding=st.session_state.encoding) > int(st.session_state.max_tokens/2):
                fragments = split_text(
                    text, 
                    max_tokens=int(st.session_state.max_tokens/2), 
                    overlap=st.session_state.overlap, 
                    encoding=st.session_state.encoding
                    )
            
            else:
                fragments = [text]
            
            # --- LLM extraction per fragment -----------------------------------
            responses_lot = []
            responses_general = []

            for fragment in fragments:
                prompt_template = f"""
                Tu es un expert en extraction d'informations pertinentes. 
                Tu es précis et tu conserves le maximum de détails. Tu donnes toujours du contexte précis.

                Mon équipe et moi nous occupons du lot suivant (il y a ici le lot et ses 
                caractéristiques, NE LES RE-MENTIONNE PAS DANS TA REPONSE) : {interesting_lots}.

                Liste moi de manière précise, exhaustive et contextualisée ce que le DOCUMENT ci-dessous 
                contient comme informations additionelles sur le lot considéré et sur 
                ses caractéristiques (trouvées dans le document UNIQUEMENT).

                Si aucune information, mets simplement 'NONE'.

                VOICI LE DOCUMENT DONT TU DOIS EXTRAIRE LES INFORMATIONS :{fragment}
                """

                prompt_template_2 = f"""
                Tu es un expert en extraction d'informations pertinentes. 
                Tu es précis et tu conserves le maximum de détails. Tu donnes toujours du contexte précis.

                Liste moi de manière précise, exhaustive et contextualisée ce que le DOCUMENT ci-dessous 
                contient comme informations générales sur le chantier, ou informations communes 
                à tous les lots (trouvées dans le document UNIQUEMENT).

                Si aucune information, mets simplement 'NONE'.

                VOICI LE DOCUMENT DONT TU DOIS EXTRAIRE LES INFORMATIONS : {fragment}
                """

                responses_lot.append(llm([HumanMessage(content=prompt_template)]).content.strip())
                responses_general.append(llm([HumanMessage(content=prompt_template_2)]).content.strip())

            combined_lot = " ".join(responses_lot)
            combined_general = " ".join(responses_general)
            summaries[file] = (
                "Informations sur le lot : " + combined_lot +
                ". Informations sur le chantier ou sur l'ensemble des lots : " + combined_general
            )
    return summaries

def markdown_to_docx(doc: DocxDocument, markdown_text: str) -> None:
    """
    Simple markdown→DOCX rendering that keeps most inline formatting
    (bold/italic, bullet lists, headings).

    The existing doc instance is modified in-place; nothing is returned.
    """ 

    html = markdown(markdown_text)
    soup = BeautifulSoup(html, "html.parser")
    
    for element in soup:
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            para = doc.add_heading("", level=level)
            add_formatted_text(para, element)

        elif element.name in ["ul", "ol"]:
            for li in element.find_all("li", recursive=False):
                para = doc.add_paragraph("", style="List Bullet")
                add_formatted_text(para, li, inline=True)
                para.paragraph_format.space_after = Pt(1)

        elif element.name == "p":
            para = doc.add_paragraph("")
            add_formatted_text(para, element)
            para.paragraph_format.line_spacing = Pt(12)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(1)

        else:
            para = doc.add_paragraph("")
            add_formatted_text(para, element)
            para.paragraph_format.line_spacing = Pt(12)

def final_summary_and_export(summaries: Dict[str, str], vectorstore):
    """
    Orchestrate the 3-pass LLM synthesis, deduplicate each section, then build
    the Word document (text + image thumbnails).

    Parameters
    ----------
    summaries : dict[str, str]
        Output of :func:`summarize_documents`.
    vectorstore : FAISS | Chroma | …
        Store used to retrieve pages and images for the picture section.

    Returns
    -------
    docx.document.Document
        Ready-to-save python-docx instance.
    """

    base_content = json.dumps(list(summaries.values()), ensure_ascii=False, indent=2)

    # ───────────────────────── PART I – chantier overview ───────────────────
    # Initial information retrieval
    prompt_part1_1 = f"""
    Tu es un expert en synthèse de documents. Tu es exhaustif et tu donnes le maximum d'informations.
    Ta tâche est UNIQUEMENT de rédiger la première partie du document final.

    Cette première partie doit présenter le chantier dans sa globalité : 
    #### Contexte général
    #### Lieux
    #### Parties prenantes et contacts

    Utilise comme contenu le prompt suivant : {base_content}
    Ne génère que cette partie et assure-toi qu'elle est bien structurée pour être intégrée dans un document Word. 
    Ne fais pas de conclusion ni d'introduction.
    """
    response1_1 = llm([HumanMessage(content=prompt_part1_1)])
    part1_1 = response1_1.content.strip()
    
    # Addition of more information
    prompt_part1_2 = f"""
    Tu es un expert en synthèse de documents. Tu es exhaustif et tu donnes le maximum d'informations.
    Ta tâche est UNIQUEMENT de compléter la première partie du document final, fournie dans le contexte.
    Contexte fourni (ne pas modifier) : {part1_1}   

    Tu dois compléter cette partie en ajoutant des informations sur le chantier en général : 
    #### Durée et échéances
    #### Différents lots disponibles, avec une description rapide
    #### Toute autre information pertinente le chantier en général

    Utilise comme contenu le prompt suivant : {base_content}

    Ne génère que cette partie et assure-toi qu'elle est bien structurée pour être intégrée dans un document Word. 
    Ne fais pas de conclusion ni d'introduction.
    """
    response1_2 = llm([HumanMessage(content=prompt_part1_2)])
    part1 = part1_1 + "\n" + response1_2.content.strip()
    
    # Delete potential duplicates
    prompt_dedup_part1 = f"""
    Tu es un expert en synthèse de documents. Ta tâche est de conserver scrupuleusement la structure 
    et l'ensemble des informations de la partie ci-dessous.
    Ta seule mission est de produire une version IDENTIQUE de cette partie en supprimant uniquement 
    les informations qui apparaissent en double. Tu conserveras l'information en double à l'endreoit le plus pertienent.
    Voici la partie à traiter : {part1}
    """
    response_dedup_part1 = llm([HumanMessage(content=prompt_dedup_part1)])
    part1 = response_dedup_part1.content.strip()
    
    # ───────────────────────── PART II – target lot ─────────────────────────
    
    # Initial information retrieval
    prompt_part2_1 = f"""
    Tu es un expert en synthèse de documents. Tu es exhaustif et tu donnes le maximum d'informations.
    Ta tâche est UNIQUEMENT d'ajouter une deuxième partie au document final dont la première partie est fournie dans le contexte.
    Contexte fourni (ne pas modifier) : {part1}

    Dans cette deuxième partie, tu dois détailler exclusivement les informations 
    relatives au lot intéressant pour l'entreprise. 

    Développe surtout sur :
    #### Une description générale du lot
    #### L'état actuel du lot
    #### Le détail des analyses réalisées au préalable sur le lot, et leurs résultats

    Utilise comme contenu le prompt suivant : {base_content}
    Ne génère que cette partie et assure-toi qu'elle s'intègre harmonieusement avec la première en faisant un tout. Ne fais de conclusion ni d'introduction.
    """
    response2 = llm([HumanMessage(content=prompt_part2_1)])
    part2_1 = response2.content.strip()
    context_1= part1 + "\n" + part2_1
    
    # Addition of more information
    prompt_part2_2 = f"""
    Tu es un expert en synthèse de documents. Tu es exhaustif et tu donnes le maximum d'informations.
    Ta tâche est UNIQUEMENT de compléter la deuxième partie du document final, cette partie est fournie dans le contexte.
    Contexte fourni (ne pas modifier) : {context_1}

    Tu dois compléter cette deuxième partie en détaillant exclusivement les informations 
    relatives au lot intéressant pour l'entreprise. 

    Développe surtout sur :
    #### Les différents travaux nécessaires pour ce lot
    #### Le matériel requis pour ce lot
    #### Les expertises et qualifications requises pour ce lot
    
    Utilise comme contenu le prompt suivant : {base_content}
    Ne génère que cette partie et assure-toi qu'elle s'intègre harmonieusement avec la première en ne faisant qu'un tout. 
    Ne fais pas de conclusion ni d'introduction.
    """
    response2_2 = llm([HumanMessage(content=prompt_part2_2)])
    part2_2 = part2_1 + "\n" + response2_2.content.strip()
    context_2 = part1 + "\n" + part2_2
    
    # Addition of more information
    prompt_part2_3 = f"""
    Tu es un expert en synthèse de documents. Tu es exhaustif et tu donnes le maximum d'informations.
    Ta tâche est UNIQUEMENT de compléter la deuxième partie du document final, cette partie est fournie dans le contexte.
    Contexte fourni (ne pas modifier) :{context_2}

    Tu dois compléter cette deuxième partie en détaillant exclusivement les informations 
    relatives au lot intéressant pour l'entreprise. 

    Développe surtout sur :
    #### Les dispositions particulières à prendre pour ce lot
    #### Toute autre information logistique ou technique spécifique à ce lot.
    Utilise comme contenu le prompt suivant : {base_content}

    Ne génère que cette partie et assure-toi qu'elle s'intègre harmonieusement avec la première 
    en en faisant qu'un tout. Ne fais pas de conclusion ni d'introduction.
    """

    response2_3 = llm([HumanMessage(content=prompt_part2_3)])
    part2 = part2_2 + "\n" + response2_3.content.strip()
    context_3 = part1 + "\n" + part2

    # Delete potential duplicates
    prompt_dedup_part2 = f"""
    Tu es un expert en synthèse de documents. 
    Ta tâche est de conserver scrupuleusement la structure et l'ensemble des informations de la partie ci-dessous.
    Ta seule mission est de produire une version IDENTIQUE de cette partie en supprimant uniquement les informations 
    qui apparaissent en double. Tu conserveras l'information en double à l'endreoit le plus pertienent.
    Voici la partie à traiter : {part2}
    """
    response_dedup_part2 = llm([HumanMessage(content=prompt_dedup_part2)])
    part2 = response_dedup_part2.content.strip()
    
    # ───────────────────────── PART III – admin info ────────────────────────
    
    # Initial information retrieval
    prompt_part3_1 = f"""
    Tu es un expert en synthèse de documents. Tu es exhaustif et tu donnes le maximum d'informations.
    Ta tâche est UNIQUEMENT d'ajouter une troisième partie au document final, fourni en contexte.
    Contexte fourni (ne pas modifier) : {context_3}

    Dans cette troisième partie, développe sur les informations administratives du chantier :
    #### Dates
    #### Prix
    #### Normes à respecter

    Utilise comme contenu le prompt suivant : {base_content}
    Ne génère que cette partie et assure-toi qu'elle s'intègre parfaitement avec les parties précédentes 
    en ne faisant qu'un tout. Ne fais pas de conclusion ni d'introduction.
    """

    response3_1 = llm([HumanMessage(content=prompt_part3_1)])
    part3_1 = response3_1.content.strip()
    context_4 = part1 + "\n" + part2 + "\n" + part3_1

    # Addition of more information
    prompt_part3_2 = f"""
    Tu es un expert en synthèse de documents. Tu es exhaustif et tu donnes le maximum d'informations.
    Ta tâche est UNIQUEMENT de compléter la troisième partie du document final, cette partie est fournie dans le contexte.
    Contexte fourni (ne pas modifier) : {context_4}
    
    Dans cette troisième partie, développe sur les informations administratives du chantier :
    #### Procédures de candidature
    #### Obligation de visite
    #### Tout autre élément administratif pertinent

    Utilise comme contenu le prompt suivant :{base_content}
    Ne génère que cette partie et assure-toi qu'elle s'intègre parfaitement avec les parties précédentes 
    en ne faisant qu'un tout. Ne fais pas de conclusion ni d'introduction.
    """

    response3_2 = llm([HumanMessage(content=prompt_part3_2)])
    part3 = part3_1 + "\n" + response3_2.content.strip()
    
    # Delete potential duplicates
    prompt_dedup_part3 = f"""
    Tu es un expert en synthèse de documents. Ta tâche est de conserver scrupuleusement la structure et l'ensemble des informations de la partie ci-dessous.
    Ta seule mission est de produire une version IDENTIQUE de cette partie en supprimant uniquement les informations qui apparaissent en double.
    Voici la partie à traiter : {part3}
    """
    response_dedup_part3 = llm([HumanMessage(content=prompt_dedup_part3)])
    part3 = response_dedup_part3.content.strip()
    
    #List unused documents
    unused_documents = [
        key for key, value in summaries.items()
        if value == "Informations sur le lot : NONE . Informations sur le chantier ou sur l'ensemble des lots : NONE"
    ]
    part4 = "\n".join(unused_documents) if unused_documents else "Aucun document non utilisé détecté."
    
    # Création of the docx document
    doc = DocxDocument()
    doc.add_heading('Synthèse Finale', 0)
    
    doc.add_heading('I. Présentation Générale du Chantier', level=1)
    markdown_to_docx(doc, part1)
    
    doc.add_heading(f'II. Informations Relatives au Lot Intéressant pour {CLIENT}', level=1)
    markdown_to_docx(doc, part2)
    
    doc.add_heading('III. Informations Administratives du Chantier', level=1)
    markdown_to_docx(doc, part3)
    
    doc.add_heading('IV. Documents non utilisés et potentiellement à explorer manuellement', level=1)
    doc.add_paragraph(part4)
    
    # --- picture section -------------------------------------------------
    images = generate_image_descriptions(f"{part1}\n{part2}\n{part3}", vectorstore)
    
    # Add a section for useful images
    if images:
        doc.add_heading('V. Images pertinentes présentes dans les documents', level=1)

        num_images = len(images)
        num_rows = (num_images + 1) // 2
        table = doc.add_table(rows=num_rows, cols=2)
        table.autofit = True
        
        for row_idx in range(num_rows):
            for col_idx in range(2):
                index = row_idx * 2 + col_idx
                cell = table.cell(row_idx, col_idx)

                if index < num_images:
                    image_info = images[index]
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    converted_image = convert_image_bytes(image_info["image_bytes"])

                    if converted_image:
                        run.add_picture(converted_image, width=Inches(2))

                    else:
                        print(f"Impossible de convertir l'image {image_info['image_name']}")
                    
                    paragraph.add_run("\n")
                    paragraph.add_run(f"Description : {image_info['description']}\n")
                    paragraph.add_run(f"Source : {image_info['image_name']}")
                
                else:
                    cell.text = ""
    
    else:
        doc.add_heading('V. Images pertinentes présentes dans les documents', level=1)
        doc.add_paragraph("Aucune image pertinente n'a été détectée.")
    
    return doc
